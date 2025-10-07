import fitz
from docx import Document
import re
from pathlib import Path
from typing import Optional, Tuple
import logging

logger = logging.getLogger(__name__)

class DocumentExtractor:

    @staticmethod
    def extract_text_from_pdf(file_path: Path) -> Tuple[str, bool]:
        try:
            doc = fitz.open(file_path)
            text_parts = []
            is_scanned = False

            for page_num, page in enumerate(doc):
                text = page.get_text("text")
                text_parts.append(text)

                if len(text.strip()) < 50 and page_num == 0:
                    is_scanned = True

            doc.close()
            full_text = "\n".join(text_parts)

            return full_text, is_scanned
        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}")
            return "", False

    @staticmethod
    def extract_text_from_docx(file_path: Path) -> str:
        try:
            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return ""

    @staticmethod
    def extract_text(file_path: Path) -> str:
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            text, is_scanned = DocumentExtractor.extract_text_from_pdf(file_path)
            if is_scanned and len(text.strip()) < 100:
                logger.warning(f"Scanned PDF detected: {file_path}. OCR not implemented.")
            return text
        elif suffix in [".docx", ".doc"]:
            return DocumentExtractor.extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file type: {suffix}")
            return ""

    @staticmethod
    def normalize_text(text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        text = text.strip()
        return text

    @staticmethod
    def extract_sections(text: str) -> dict:
        sections = {
            "education": "",
            "experience": "",
            "skills": "",
            "projects": "",
            "certifications": "",
            "publications": "",
            "full_text": text
        }

        education_patterns = [
            r"(?i)(education|academic background|qualifications)(.*?)(?=(experience|employment|skills|projects|certifications|$))",
            r"(?i)(university|college|bachelor|master|phd|mba)(.*?)(?=(experience|employment|skills|projects|$))"
        ]

        experience_patterns = [
            r"(?i)(experience|employment|work history|professional experience)(.*?)(?=(education|skills|projects|certifications|publications|$))",
            r"(?i)(analyst|associate|intern|engineer|manager|director)(.*?)(?=(education|skills|projects|$))"
        ]

        skills_patterns = [
            r"(?i)(skills|technical skills|competencies|expertise)(.*?)(?=(experience|education|projects|certifications|$))",
            r"(?i)(python|java|c\+\+|sql|machine learning)(.*?)(?=(experience|education|projects|$))"
        ]

        for pattern in education_patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match and len(match.group(0)) > len(sections["education"]):
                sections["education"] = match.group(0)

        for pattern in experience_patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match and len(match.group(0)) > len(sections["experience"]):
                sections["experience"] = match.group(0)

        for pattern in skills_patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match and len(match.group(0)) > len(sections["skills"]):
                sections["skills"] = match.group(0)

        return sections

class PIIExtractor:

    EMAIL_PATTERN = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+(?:\.[A-Z|a-z]{2,})?\b')
    PHONE_PATTERNS = [
        re.compile(r'\+\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}'),
        re.compile(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'),
        re.compile(r'\d{3}[-.\s]\d{3}[-.\s]\d{4}'),
        re.compile(r'\d{10,15}')
    ]
    LINKEDIN_PATTERN = re.compile(r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+/?', re.IGNORECASE)

    @staticmethod
    def extract_emails(text: str) -> list:
        return list(set(PIIExtractor.EMAIL_PATTERN.findall(text)))

    @staticmethod
    def validate_email(email: str) -> bool:
        if not email:
            return False
        full_email_pattern = r'^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$'
        return bool(re.match(full_email_pattern, email))

    @staticmethod
    def extract_phones(text: str) -> list:
        phones = set()
        for pattern in PIIExtractor.PHONE_PATTERNS:
            matches = pattern.findall(text)
            for match in matches:
                digits = re.sub(r'\D', '', match)
                if 10 <= len(digits) <= 15:
                    phones.add(match.strip())
        return list(phones)

    @staticmethod
    def extract_linkedin(text: str) -> Optional[str]:
        match = PIIExtractor.LINKEDIN_PATTERN.search(text)
        return match.group(0) if match else None

    @staticmethod
    def mask_pii(text: str) -> str:
        text = PIIExtractor.EMAIL_PATTERN.sub('[EMAIL]', text)
        for pattern in PIIExtractor.PHONE_PATTERNS:
            text = pattern.sub('[PHONE]', text)
        text = PIIExtractor.LINKEDIN_PATTERN.sub('[LINKEDIN]', text)
        return text

    @staticmethod
    def extract_all_pii(text: str) -> dict:
        return {
            "emails": PIIExtractor.extract_emails(text),
            "phones": PIIExtractor.extract_phones(text),
            "linkedin": PIIExtractor.extract_linkedin(text)
        }
